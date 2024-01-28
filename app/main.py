import os
from openpyxl import load_workbook
from orquesta_sdk import Orquesta, OrquestaClientOptions
from docx import Document
from dotenv import load_dotenv
from fastapi import FastAPI, HTTPException, Body

load_dotenv()

# Initialize Orquesta client
def init_orquesta_client():
    api_key = os.getenv("ORQUESTA_API_KEY")
    options = OrquestaClientOptions(api_key=api_key, environment="production")
    return Orquesta(options)

client = init_orquesta_client()

app = FastAPI()

# Load questions from an Excel file
def load_questions_from_sheet(sheet_path):
    questions = []
    workbook = load_workbook(sheet_path)
    sheet = workbook.active
    for row in sheet.iter_rows(min_row=2, values_only=True):  # Skip header row
        question = row[2]  # Question is in third column
        if question:  # Check if the cell is not empty
            questions.append(question)
    return questions

# Define the path for the output Word document
output_doc_path = 'survey_results.docx'

# Initialize the Word document
document = Document()

# Initialize variables to store the previous question and response
previous_question = ""
previous_response = ""

# Path to your Excel file containing the questions
sheet_path = 'app/vragenlijsten-V3.xlsx'
questions = load_questions_from_sheet(sheet_path)

@app.post("/chat/")
async def chat(user_message: str = Body(..., embed=True)):
    global previous_question, previous_response, document
    try:
        # Find the next question to ask
        if not previous_question:
            question_to_ask = questions[0]
        else:
            question_to_ask = questions[questions.index(previous_question) + 1]

        # Combine the previous question and response
        previous_combined = "Vorige vraag: " + f"{previous_question}" + " Antwoord: " + f"{previous_response}" if previous_question else ""

        # Send the question to the deployment for rephrasing
        deployment = client.deployments.invoke(
            key="logopedica-vragenlijsten",
            context={
                "environments": [],
                "klacht": ["mondgewoonten"]
            },
            inputs={
                "question": question_to_ask,
                "previous": previous_combined.strip()  # Remove leading/trailing whitespace
            }
        )

        # Get the rephrased question from the deployment
        rephrased_question = deployment.choices[0].message.content

        # Add the user's message and the rephrased question to the Word document
        document.add_paragraph(f"Vraag: {previous_question}")
        document.add_paragraph(f"Antwoord: {user_message}")
        document.add_paragraph("")  # Add a blank line between Q&A pairs

        # Save the Word document after each Q&A pair
        document.save(output_doc_path)

        # Update the previous question and response
        previous_question = question_to_ask
        previous_response = user_message

        # Return the rephrased question to the front-end
        return {"response": rephrased_question}

    except IndexError:
        # No more questions to ask, end the survey
        return {"response": "Survey completed. Thank you for your participation."}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))